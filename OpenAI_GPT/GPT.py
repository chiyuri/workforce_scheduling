# ------------------Copyright (C) 2023 University of Strathclyde and Author ---------------------------------
# --------------------------------- Author: Cheyenne Powell -------------------------------------------------
# ------------------------- e-mail: cheyenne.powell@strath.ac.uk --------------------------------------------

# Loads the schedule data into openAI GPT3.5 with constraints provided to see how it responds due to conditions
# ===========================================================================================================

import sys
import os
import openai
import pandas as pd
from Extract_solution import Extract_solution
from Extract_input import Extract_input
import numpy

import docx

# open connection to Word Document
doc = docx.Document("../Questions_answers.docx")
p = doc.add_paragraph()
p.add_run("Question:").bold = True

Question = input("please enter question here: \n")
doc.add_paragraph(Question)
# print(Question)

# result = [p.text for p in doc.paragraphs]
# print(result)


#\n Mary is sick on Tuesday and is unable to work. Is there anyone able to cover her shift?\n"
def OpenAI_connect():
    openai.api_key = '<enter code here>'
    # Loading the solvers results.
    data_path1 = "../data/InputData.xlsx"
    data_path2 = '../Solution.xlsx'

    data_employees, demand, parameters, input_days, optimization_parameters = Extract_input(data_path1)
    daily_hrs_employees, weeklyHours = Extract_solution(data_path2)
    # emp_hrs_sched = list(daily_hrs_employees.index)
    # emp_wk_hrs = list(weeklyHours)



    data = pd.read_excel(r''+data_path2)
    df = pd.DataFrame(data, columns=['Employee', 'Date', 'Day', 'Start', 'End', 'Duration_hrs'])
    # df2 = pd.DataFrame(data_employees, columns=['Employee', 'Date', 'Day', 'Start', 'End', 'Duration_hrs'])
    gpt_prompt = str(df)
    # print(str(data), '\n',str(df), '\n', str(data_employees))


    # print(gpt_prompt)


    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[{"role": "user", "content": "This is a workforce schedule with constraints.\n"},
        {"role": "system", "content": "The data below contains the number of hours each employee is available\n"},
        {"role": "system", "content": str(data_employees)},
        {"role": "system", "content": "\n\nThis is the demand for special skills from each employee\n"},
        {"role": "system", "content": str(demand)},
        {"role": "system", "content": "\nThe minimum working time per day is 4 hrs"},
        {"role": "system", "content": "\nThe maximum working time per day is 8 hrs"},
        {"role": "system", "content": "\nThe demand of special qualification is 2 per slot"},
        {"role": "system", "content": "\nThe maximum hours per week varies with each person, using the availability"},
        {"role": "system", "content": "\nThe minimum number of hrs per week varies with each person\n"},
        {"role": "system", "content": "\nThere is no overtime"},
        {"role": "system", "content": "These are the dates:\n"},
        {"role": "system", "content": str(input_days)},
        {"role": "system", "content": "\nThese are the hrs each employee is scheduled to work\n"},
        {"role": "system", "content": str(daily_hrs_employees)},
        {"role": "system", "content": "\nQuestion:\n"},
        {"role": "system", "content": Question},
        { "role": "user", "content": ""}],
    #       {"role": "user", "content": "\n Who is not working on Monday?\n"}],
        temperature=1.13,
        max_tokens=256,
        top_p=1,
        frequency_penalty=0,
        presence_penalty=0
    )
    # print(response)
    # response.choices[0].text
    # print(response.choices[0].text)
    numpy.set_printoptions(threshold=sys.maxsize)
    print(response['choices'][0]["message"]["content"])
    p = doc.add_paragraph()
    p.add_run("Answer:").bold = True
    doc.add_paragraph(response['choices'][0]["message"]["content"])
    doc.save("../Questions_answers.docx")
OpenAI_connect()